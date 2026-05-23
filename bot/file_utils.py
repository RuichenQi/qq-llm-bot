"""Read text out of common file formats so the bot can answer questions about
attachments the way it does about images.

Each extractor is best-effort: optional deps (pypdf, python-docx) are imported
lazily; missing deps degrade to a "I can't read this format" error rather than
crashing the bot.

Audio / video are handled out-of-band by the caller (audio → Whisper, video
→ ffmpeg frame sampling + Whisper), since they need provider access. This
module only classifies them so the caller knows which path to take.
"""
from __future__ import annotations

import io
import json
import re
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple


# Tag groups → handler kind.
_KIND_BY_EXT = {
    # plain text
    ".txt": "text", ".md": "text", ".log": "text", ".csv": "text",
    ".tsv": "text", ".json": "text", ".yaml": "text", ".yml": "text",
    ".xml": "text", ".html": "text", ".htm": "text", ".srt": "text",
    ".ini": "text", ".toml": "text", ".env": "text", ".rst": "text",
    # code (treated as text but rendered as code-block in the prompt)
    ".py": "code", ".js": "code", ".ts": "code", ".tsx": "code", ".jsx": "code",
    ".go": "code", ".rs": "code", ".java": "code", ".kt": "code",
    ".c": "code", ".h": "code", ".cpp": "code", ".hpp": "code", ".cc": "code",
    ".cs": "code", ".rb": "code", ".php": "code", ".swift": "code",
    ".sh": "code", ".bash": "code", ".zsh": "code", ".ps1": "code", ".bat": "code",
    ".sql": "code", ".vue": "code", ".svelte": "code", ".css": "code", ".scss": "code",
    # documents
    ".pdf": "pdf",
    ".docx": "docx",
    # audio
    ".mp3": "audio", ".wav": "audio", ".m4a": "audio", ".ogg": "audio",
    ".flac": "audio", ".aac": "audio", ".opus": "audio", ".webm": "audio",
    ".amr": "audio",
    # video
    ".mp4": "video", ".mov": "video", ".mkv": "video", ".avi": "video",
    ".flv": "video", ".wmv": "video",
}


@dataclass
class FileExtraction:
    kind: str           # text | code | pdf | docx | audio | video | unsupported
    name: str
    text: str           # extracted text (may be empty for audio/video → caller handles)
    truncated: bool = False
    error: Optional[str] = None


def classify(name: str) -> str:
    ext = Path(name).suffix.lower()
    return _KIND_BY_EXT.get(ext, "unsupported")


def _decode_text(data: bytes) -> str:
    """Best-effort decode: try UTF-8 then GBK / fall back to latin-1."""
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030", "big5"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("latin-1", errors="replace")


def extract_text(name: str, data: bytes, *, max_chars: int) -> FileExtraction:
    """Pull plain text out of `data` given its filename. Caller is responsible
    for size limits BEFORE calling — we don't sniff content."""
    kind = classify(name)
    if kind == "text":
        text = _decode_text(data)
        return _truncate(FileExtraction(kind="text", name=name, text=text), max_chars)
    if kind == "code":
        text = _decode_text(data)
        ext = Path(name).suffix.lstrip(".")
        text = f"```{ext}\n{text}\n```"
        return _truncate(FileExtraction(kind="code", name=name, text=text), max_chars)
    if kind == "pdf":
        return _truncate(_extract_pdf(name, data), max_chars)
    if kind == "docx":
        return _truncate(_extract_docx(name, data), max_chars)
    if kind in ("audio", "video"):
        # Caller handles via provider (Whisper / frame sampling).
        return FileExtraction(kind=kind, name=name, text="")
    return FileExtraction(
        kind="unsupported", name=name, text="",
        error=f"暂不支持的文件类型: {Path(name).suffix or '(无扩展名)'}",
    )


def _truncate(extr: FileExtraction, max_chars: int) -> FileExtraction:
    if extr.text and len(extr.text) > max_chars:
        extr.text = extr.text[:max_chars]
        extr.truncated = True
    return extr


def _extract_pdf(name: str, data: bytes) -> FileExtraction:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except ImportError:
            return FileExtraction(
                kind="pdf", name=name, text="",
                error="读 PDF 需要 pypdf 库（pip install pypdf）",
            )
    try:
        reader = PdfReader(io.BytesIO(data))
        parts: List[str] = []
        for i, page in enumerate(reader.pages):
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                parts.append("")
            if i >= 200:  # hard page cap
                break
        text = "\n".join(p.strip() for p in parts if p.strip())
        if not text.strip():
            return FileExtraction(
                kind="pdf", name=name, text="",
                error="这个 PDF 没抽出文字（可能是扫描件 / 图片）",
            )
        return FileExtraction(kind="pdf", name=name, text=text)
    except Exception as e:
        return FileExtraction(
            kind="pdf", name=name, text="",
            error=f"PDF 读取失败: {e!r}",
        )


def _extract_docx(name: str, data: bytes) -> FileExtraction:
    try:
        import docx  # python-docx  # type: ignore
    except ImportError:
        return FileExtraction(
            kind="docx", name=name, text="",
            error="读 Word 文档需要 python-docx 库（pip install python-docx）",
        )
    try:
        doc = docx.Document(io.BytesIO(data))
        parts: List[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also pull table cells, in row-major order.
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        if not text.strip():
            return FileExtraction(
                kind="docx", name=name, text="",
                error="这个 docx 是空的",
            )
        return FileExtraction(kind="docx", name=name, text=text)
    except Exception as e:
        return FileExtraction(
            kind="docx", name=name, text="",
            error=f"docx 读取失败: {e!r}",
        )


# ---------- ffmpeg helpers ----------
def _ffmpeg_available() -> bool:
    try:
        subprocess.run(
            ["ffmpeg", "-version"],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
            check=True, timeout=4,
        )
        return True
    except (OSError, subprocess.SubprocessError):
        return False


def downsample_audio(data: bytes, ext: str) -> Optional[bytes]:
    """Re-encode audio to 16kHz mono MP3 to slash Whisper upload size.
    Returns the new bytes or None if ffmpeg is missing / fails."""
    if not _ffmpeg_available():
        return None
    with tempfile.TemporaryDirectory() as tmp:
        in_path = Path(tmp) / f"in{ext or '.bin'}"
        out_path = Path(tmp) / "out.mp3"
        in_path.write_bytes(data)
        try:
            subprocess.run(
                [
                    "ffmpeg", "-y", "-i", str(in_path),
                    "-ac", "1", "-ar", "16000", "-b:a", "32k",
                    str(out_path),
                ],
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                check=True, timeout=120,
            )
        except (OSError, subprocess.SubprocessError):
            return None
        if not out_path.exists():
            return None
        return out_path.read_bytes()


def extract_video_frames(data: bytes, ext: str, count: int) -> List[bytes]:
    """Extract `count` evenly-spaced PNG frames. Empty list if ffmpeg missing."""
    if count <= 0 or not _ffmpeg_available():
        return []
    with tempfile.TemporaryDirectory() as tmp:
        in_path = Path(tmp) / f"in{ext or '.mp4'}"
        in_path.write_bytes(data)
        # Get duration via ffprobe; fall back to a fixed 4s gap if unavailable.
        try:
            res = subprocess.run(
                [
                    "ffprobe", "-v", "error", "-show_entries", "format=duration",
                    "-of", "json", str(in_path),
                ],
                capture_output=True, check=True, timeout=10,
            )
            dur = float(json.loads(res.stdout.decode()).get("format", {})
                        .get("duration", 0.0))
        except Exception:
            dur = 0.0
        frames: List[bytes] = []
        for i in range(count):
            # Evenly space: at 5%, 30%, 55%, 80% of duration (or fixed intervals).
            offset = (dur * (i + 0.5) / count) if dur > 0 else (i + 1) * 2.0
            out_path = Path(tmp) / f"f{i}.png"
            try:
                subprocess.run(
                    [
                        "ffmpeg", "-y", "-ss", f"{offset:.2f}",
                        "-i", str(in_path), "-vframes", "1",
                        "-vf", "scale='min(512,iw)':-2",
                        str(out_path),
                    ],
                    stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                    check=True, timeout=20,
                )
            except (OSError, subprocess.SubprocessError):
                continue
            if out_path.exists():
                frames.append(out_path.read_bytes())
        return frames


# ---------- prompt formatting ----------
def format_for_prompt(extr: FileExtraction) -> str:
    """Render an extraction as a system-prompt block."""
    header = f"[用户上传的文件: {extr.name}]"
    if extr.error:
        return f"{header}\n（{extr.error}）"
    if not extr.text:
        return f"{header}\n（内容为空）"
    body = extr.text
    if extr.truncated:
        body += "\n…（内容已截断）"
    return f"{header}\n{body}"


# Stripped filename for safety / logging.
_SAFE_NAME_RE = re.compile(r"[^\w.\-]+", re.UNICODE)


def safe_name(name: str) -> str:
    return _SAFE_NAME_RE.sub("_", (name or "file"))[:80] or "file"
